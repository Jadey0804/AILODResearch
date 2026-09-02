from copy import deepcopy
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(r"C:\WarwickProjects\AILODResearch")
SOURCE = PROJECT_ROOT / "Deliverables" / "ProjectGuide" / "AILOD_Project_Iteration_and_Runtime_Guide_CN_v1.1_Beginner.docx"
OUTPUT = PROJECT_ROOT / "Deliverables" / "ProjectGuide" / "AILOD_Project_Iteration_and_Runtime_Guide_CN_v1.2_FinalData.docx"


def set_paragraph(document: Document, index: int, text: str) -> None:
    paragraph = document.paragraphs[index]
    paragraph.text = text


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def copy_row_format(source_row, target_row) -> None:
    target_row._tr.get_or_add_trPr()
    if source_row._tr.trPr is not None:
        target_row._tr.remove(target_row._tr.trPr)
        target_row._tr.insert(0, deepcopy(source_row._tr.trPr))
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        if target_cell._tc.tcPr is not None:
            target_cell._tc.remove(target_cell._tc.tcPr)
        target_cell._tc.insert(0, deepcopy(source_cell._tc.tcPr))


def append_table_row(table, values: list[str]) -> None:
    source_row = table.rows[-1]
    row = table.add_row()
    copy_row_format(source_row, row)
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)


def main() -> None:
    document = Document(SOURCE)
    document.core_properties.title = "AILOD 项目零基础理解手册 v1.2"
    document.core_properties.subject = "大规模 NPC Simulation LOD：项目迭代、运行管线和正式结果"

    set_paragraph(document, 7, "版本快照：phase-8-formal-experiments · 7699bbf")
    set_paragraph(document, 8, "状态：Phase 7F 有画面工程验收和 Phase 8 正式 480/90 次实验均已完成；项目进入论文、PDF、视频与版本封存阶段")
    set_paragraph(document, 9, "更新日期：2026-09-01")
    set_paragraph(document, 21, "仓库依据：Git HEAD 7699bbf；Docs/AILOD_MVP_Current_Rules_Index_CN.md；Phase 7F 与 Phase 8 最终检查点；正式实验原始归档哈希。")

    set_paragraph(document, 66, "证据纪律  NullRHI、Development、单个 Seed 和工程压力结果仍只用于工程诊断。论文统计使用冻结后的 480/90 次 Shipping 正式运行；它们通过资格审计，并已用 paired bootstrap 与逐 Seed 结果分析。")

    set_paragraph(document, 172, "B5E 的 5.226 倍仍只属于工程先导结果。Phase 8 的正式 Shipping 数据给出：2k 配对中位速度比 0.140255，10k 为 1.295480，20k 为 3.943749；正式结论以 90 次运行、执行顺序敏感性和区间共同解释。")

    set_paragraph(document, 186, "当前阶段判断  模拟、实验工具、可视化、有画面工程证据和正式研究数据已经形成闭环。接下来只做论文回填、最终 PDF、答辩视频和版本封存，不再扩展研究范围。")

    set_paragraph(document, 204, "9.2 Phase 8 正式性能与准确性证据")
    set_paragraph(document, 205, "冻结提交 7699bbf 共完成 480/480 次 200 人准确性运行和 90/90 次 2k/10k/20k 性能运行，全部为 Shipping（发布优化构建）正式数据，四组 hard error（会直接使运行失去资格的硬错误）总数均为 0。20k 时 Proposed 相对 Per-Agent 的配对中位加速为 3.943749 倍；10k 开始占优；2k 因固定管理成本而更慢。")
    set_paragraph(document, 206, "准确性结果  Proposed 的 Behavior TVD（行为比例分布差异）四个场景中位数约为 0.0003–0.0004，Simple 约为 0.775–0.785。Proposed 保持 ResidentID、HomeID、住房状态、目标和第一步行为；普通 Routine 的 TaskActive（任务是否仍在进行）时间相位与 Oracle 明显不一致。这条负面结果必须和性能收益一起报告。")
    set_paragraph(document, 209, "把四类证据分开记：编译通过说明程序能生成；自动回归说明已编码规则没有被新改动破坏；Phase 7F 有画面矩阵说明当前 Demo 在指定机器与路线下怎样运行；Phase 8 的 480/90 次 Shipping 正式实验、paired bootstrap（按同一随机种子成对重采样得到置信区间）和原始哈希才支持论文研究结论。")

    set_paragraph(document, 229, "12. 项目最后还要完成什么")
    set_paragraph(document, 230, "12.1 Phase 7F：有画面工程结果已经完成")
    set_paragraph(document, 231, "Phase 7F 完成 24 格 Development（开发构建）有画面矩阵，并在 NavMesh fixed tile pool（导航网格固定瓦片池）扩大和重建后复测 4 个 World Partition（地图分块加载）配置。普通镜头和 44 个完整 Actor 绑定较稳定；20k/4x 的 World Partition 往返仍出现明显 Game Thread（游戏主线程）长帧。该证据描述 Demo 画面成本，不替代 Phase 8 的 NullRHI 后台算法统计。")
    set_paragraph(document, 232, "12.2 Phase 8：正式实验和统计已经完成")
    set_paragraph(document, 233, "正式数据  Accuracy（准确性）480/480 Runs；Performance（性能）2k、10k、20k 各 30/30 Runs，共 90/90。所有 Run 使用同一冻结 Git 提交、Shipping 构建和资格审计。")
    set_paragraph(document, 234, "统计方法  准确性按同一 Scenario + Metric + Seed 配对，使用 100,000 次 deterministic paired bootstrap（确定性成对自助重采样）报告均值差 95% CI；性能按完整 67 天 AI CPU 总时间和 10 次重复报告中位数、IQR、P95 与区间。")
    set_paragraph(document, 235, "性能发现  Proposed 在 2k 的固定成本使它慢于 Per-Agent；10k 的配对中位速度比为 1.295480；20k 为 3.943749。Simple 仍最快，它的宏观和行为误差明显更大。")
    set_paragraph(document, 236, "连续性发现  TaskActive 在四个场景的 30/30 Seeds 都越过预冻结 10% 审查线。每场景 1,800 对快照中有 1,603 对不一致，全部来自 Proposed=Active、Oracle=Inactive，且双方仍有相同 Goal、FirstAction、HomeState 和 Routine 类型；RestoreHome 等承诺事件相关不一致为 0。")
    set_paragraph(document, 237, "可信边界  Proposed 的进程峰值内存高于 Per-Agent，现有证据不支持总内存节省结论；NullRHI 结果不能改写成 FPS；单机结果不能代表所有硬件和地图；执行位置与耗时仍有一定相关，长批次热状态和时间漂移需要写入限制。")
    set_paragraph(document, 238, "12.3 现在真正剩下的工作")
    set_paragraph(document, 239, "用户完成论文修改后，再一次性填入正式数字并审核 Abstract、Results、Discussion 和 Conclusion；随后从 Word 导出最终 PDF 并逐页检查。最后由用户录制 2–4 分钟答辩视频。得到明确批准后，再提交本轮材料、创建版本标签，并 push 或制作 git bundle（完整 Git 离线备份包）。")

    set_paragraph(document, 297, "□ 我知道 100,000 人压力运行、Phase 7F 有画面矩阵、73/73 最终回归和 Phase 8 的 480/90 正式实验分别能证明什么，也知道它们不能互相替代。")
    set_paragraph(document, 298, "□ 我能说出正式结果的四个核心事实：2k 固定成本劣势、20k 约 3.94 倍性能优势、TaskActive 时间相位负面发现、没有进程峰值内存优势。")
    set_paragraph(document, 299, "□ 我能诚实说明准星/遮挡、World Partition 长帧、身份静态内存 O(N)、TaskActive 语义差异、单机 NullRHI 和无玩家研究等限制。")

    current_table = document.tables[2]
    set_cell_text(current_table.cell(5, 0), "冻结正式提交")
    set_cell_text(current_table.cell(5, 1), "7699bbf；Phase 8 正式运行均记录该提交和 Shipping 构建")
    set_cell_text(current_table.cell(6, 0), "当前证据")
    set_cell_text(current_table.cell(6, 1), "最终自动回归 73/73；Phase 7F 有画面 24 格 + NavMesh 修正 4 格；Phase 8 准确性 480/480、性能 90/90；hard error=0")
    set_cell_text(current_table.cell(7, 0), "仍待完成")
    set_cell_text(current_table.cell(7, 1), "论文真实数据回填与全文审核、最终 PDF 逐页检查、用户录制答辩视频、批准后的提交/标签/备份")

    evidence_table = document.tables[20]
    set_cell_text(evidence_table.cell(4, 1), "73/73 通过")
    append_table_row(evidence_table, ["Phase 7F 有画面矩阵", "24 格 + NavMesh 修正 4 格", "当前 Demo 在指定机器、镜头、人口和倍率下的帧时间与加载尖峰"])
    append_table_row(evidence_table, ["Phase 8 正式实验", "Accuracy 480/480；Performance 90/90；hard error=0", "正式性能、准确性、连续性和限制，可用于论文研究结论"])

    risk_table = document.tables[22]
    set_cell_text(risk_table.cell(2, 0), "Simple 的速度可能掩盖误差")
    set_cell_text(risk_table.cell(2, 1), "正式数据表明 Simple 极快，但宏观、政策和 Behavior TVD 误差明显大于 Proposed；复杂架构的价值来自精度与性能折中。")
    set_cell_text(risk_table.cell(10, 0), "正式统计的负面发现")
    set_cell_text(risk_table.cell(10, 1), "480/90 正式运行已完成。TaskActive 在四个场景的 30/30 Seeds 越线，且 Proposed 的进程峰值内存没有优势；论文必须正面报告。")

    claims_table = document.tables[24]
    set_cell_text(claims_table.cell(0, 0), "可以写成已验证事实")
    set_cell_text(claims_table.cell(0, 1), "必须同时写清的边界")
    set_cell_text(claims_table.cell(1, 0), "v1.9 架构、单一权威、Batch、Lift/Restrict、住房连续性和 Demo 只读边界已经实现并验证")
    set_cell_text(claims_table.cell(1, 1), "Proposed 与 Oracle 存在受控近似；普通 Routine 的 TaskActive 时间相位明显不一致")
    set_cell_text(claims_table.cell(2, 0), "20k 时 Proposed 相对 Per-Agent 的配对中位加速约 3.94 倍；10k 开始占优")
    set_cell_text(claims_table.cell(2, 1), "2k 时 Proposed 更慢；结果限定于冻结的单机、Shipping、NullRHI 和当前场景")
    set_cell_text(claims_table.cell(3, 0), "Proposed 的宏观、政策和 Behavior TVD 明显优于 Simple；Phase 7F 已记录有画面表现")
    set_cell_text(claims_table.cell(3, 1), "进程峰值内存没有 Proposed 优势；准星/遮挡和 World Partition 长帧继续作为 Demo 限制")

    timeline = document.tables[25]
    append_table_row(timeline, ["2026-08-31", "3314034", "封板 Phase 7F 有画面矩阵并修正 NavMesh 容量"])
    append_table_row(timeline, ["2026-08-31", "cb8a39b", "关闭 Phase 8 正式实验管线"])
    append_table_row(timeline, ["2026-08-31", "7699bbf", "解除 Shipping 实验构建阻塞并冻结正式实验提交"])

    footer = document.sections[0].footer.paragraphs[0]
    if footer.runs:
        footer.runs[0].text = "零基础正式数据修订版 v1.2 · 7699bbf   |   第 "

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
